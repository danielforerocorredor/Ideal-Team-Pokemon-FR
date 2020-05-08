import networkx as nx
import matplotlib.pyplot as plt
import matplotlib._color_data as mcd
import matplotlib.patches as patches
from matplotlib.offsetbox import AnnotationBbox
import random
from docx import Document
from docx.shared import Inches

#Creacion clase Pokemon
class pokemon:
  def __init__(self, name, ptype, site, rate):
    self.name = name
    self.ptype = ptype
    self.site = site
    self.rate = rate

#Creacion clase Ciudad.
#Sirve para saber si la ciudad tiene gymnasio o no, de tener, poder saber el tipo
class Ciudad:
  def __init__(self, name, pokeType, gym, leader):
    self.name = name
    self.pokeType = pokeType
    self.gym = gym
    self.leader = leader

#clase que representa los tipos de pokemon odenados de mayor a menor.
class Tipos:
  def __init__(self, name, debilidades):
    self.name = name
    self.debilidades = debilidades
  #funcion retorna debilidad de un tipo especifico
  def getWeaknessType(self):
    myRand = random.randint(0,len(self.debilidades)-1)
    return(self.debilidades[myRand])


#Creacion tipos
debilidadesRoca = ['Agua', 'Hierba', 'Lucha', 'Electrico']
Roca = Tipos('Roca', debilidadesRoca)

debilidadesBicho = ['fuego', 'Roca', 'volador', 'Normal']
Bicho = Tipos('Bicho', debilidadesBicho)

debilidadesPsiquico = ['Bicho', 'Fantasma', 'Normal', 'Fuego']
Psiquico = Tipos('Psiquico', debilidadesPsiquico)

debilidadesVolador = ['Roca', 'Electrico', 'Hielo']
Volador = Tipos('Volador', debilidadesVolador)

debilidadesTierra = ['Agua', 'Hierba', 'Hielo']
Tierra = Tipos('Tierra', debilidadesTierra)

debilidadesVeneno = ['Psiquico', 'Tierra', 'Hielo']
Veneno = Tipos('Veneno', debilidadesVeneno)

debilidadesLucha = ['Volador', 'Psiquico', 'Hierba']
Lucha = Tipos('Lucha', debilidadesLucha)

debilidadesHielo = ['Fuego', 'Lucha', 'Roca']
Hielo = Tipos('Hielo', debilidadesHielo)

debilidadesHierba = ['Fuego', 'Hielo', 'Veneno', 'Volador', 'Bicho']
Hierba = Tipos('Hierba', debilidadesHierba)

debilidadesElectrico = ['Tierra', 'Normal']
Electrico = Tipos('Electrico', debilidadesElectrico)

debilidadesAgua = ['Electrico', 'Hierba']
Agua = Tipos('Agua', debilidadesAgua)

debilidadesFuego = ['Agua', 'Tierra', 'Roca']
Fuego = Tipos('Fuego', debilidadesFuego)

DebilidadesNormal = ['Lucha', 'Roca']
Normal = Tipos('Normal', DebilidadesNormal)

debilidadesFantasma = ['Fuego', 'Agua', 'Electrico', 'Hierba', 'Dragon', 'Volador']
Fantasma = Tipos('Fantasma', debilidadesFantasma)

debilidadesDragon = ['Hielo']
Dragon = Tipos('Dragon', debilidadesDragon)

#Creacion ciudades
    #1. tienen gimnasio
CiudadPlateada = Ciudad('Ciudad Plateada', Roca, True, 'Brock')
CiudadCeleste = Ciudad('Ciudad Celeste', Agua, True, 'Misty')
CiudadCarmin = Ciudad('Ciudad Carmín', Electrico, True, 'Lt. Surge')
CiudadAzulona = Ciudad('Ciudad Azulona', Hierba, True, 'Erika')
CiudadFucsia = Ciudad('Ciudad Fucsia', Veneno, True, 'Koga')
CiudadAzafran = Ciudad('Ciudad Azafrán', Psiquico, True, 'Sabrina')
IslaCanela = Ciudad('Isla Canela', Fuego, True, 'Blaine')
CiudadVerde = Ciudad('Ciudad Verde', Tierra, True, 'Giovanni')
    #2. no tienen gimnasio
MesetaAnil = Ciudad('Meseta Añil', '', False, '')
CalleVictoria = Ciudad('Calle Victoria', '', False, '')
PuebloPaleta = Ciudad('Pueblo Paleta', '', False, '')
BosqueVerde = Ciudad('Bosque Verde', '', False, '')
CuevaDiglet = Ciudad('Cueva Diglet', '', False, '')
MonteLuna = Ciudad('Monte Luna', '', False, '')
TunelRoca = Ciudad('Tunel Roca', '', False, '')
CentralEnergia = Ciudad('Central Energía', '', False, '')
PuebloLavanda = Ciudad('Pueblo Lavanda', '', False, '')
IslasEspuma = Ciudad('Islas Espuma', '', False, '')

#Lista Ciudades

listaCiudades = []
listaCiudades.append(CiudadPlateada)
listaCiudades.append(CiudadCeleste)
listaCiudades.append(CiudadCarmin)
listaCiudades.append(CiudadAzulona)
listaCiudades.append(CiudadFucsia)
listaCiudades.append(CiudadAzafran)
listaCiudades.append(IslaCanela)
listaCiudades.append(CiudadVerde)
#
listaCiudades.append(MesetaAnil)
listaCiudades.append(CalleVictoria)
listaCiudades.append(PuebloPaleta)
listaCiudades.append(BosqueVerde)
listaCiudades.append(CuevaDiglet)
listaCiudades.append(MonteLuna)
listaCiudades.append(TunelRoca)
listaCiudades.append(CentralEnergia)
listaCiudades.append(PuebloLavanda)
listaCiudades.append(IslasEspuma)

#Lista en orden de paradas para acar el juego

listaCiudadesEnOrden = []
listaCiudadesEnOrden.append(PuebloPaleta)#1
listaCiudadesEnOrden.append(CiudadVerde)#2
listaCiudadesEnOrden.append(BosqueVerde)#3
listaCiudadesEnOrden.append(CuevaDiglet)#4
listaCiudadesEnOrden.append(CiudadPlateada)#5
listaCiudadesEnOrden.append(MonteLuna)#6
listaCiudadesEnOrden.append(CiudadCeleste)#7
listaCiudadesEnOrden.append(CiudadAzafran)#8
listaCiudadesEnOrden.append(CiudadCarmin)#9
listaCiudadesEnOrden.append(CiudadAzafran)#10
listaCiudadesEnOrden.append(CiudadAzulona)#11
listaCiudadesEnOrden.append(CiudadFucsia)#12
listaCiudadesEnOrden.append(CiudadCarmin)#13
listaCiudadesEnOrden.append(CiudadAzafran)#14
listaCiudadesEnOrden.append(PuebloLavanda)#15
listaCiudadesEnOrden.append(CiudadFucsia)#16
listaCiudadesEnOrden.append(IslasEspuma)#17
listaCiudadesEnOrden.append(IslaCanela)#18
listaCiudadesEnOrden.append(PuebloPaleta)#19
listaCiudadesEnOrden.append(CiudadVerde)#20


#Lista de gimnasios en orden
listaGimnasiosEnOrden = []
for i in listaCiudades:
  if i.gym == True:
    if i not in listaGimnasiosEnOrden:
      listaGimnasiosEnOrden.append(i)


#funcion que retorna la debilida de una ciudad
def getWeaknessTypeCity(city):
  return city.pokeType.getWeaknessType()

#for i in listaGimnasiosEnOrden:
#  print('la debilidad de',i.name,'cuyo tipo es ',i.pokeType.name,' es',getWeaknessTypeCity(i))

#creacion pokemon dependiendo de en donde se encuenrtan
myPokes = []
myStarters = []

#             STARTERS
Bulbasaur = pokemon('Bulbasaur', Hierba, PuebloPaleta, 0)
myStarters.append(Bulbasaur)
Charmander = pokemon('Charmander', Fuego, PuebloPaleta, 0)
myStarters.append(Charmander)
Squirtle = pokemon('Squirtle', Agua, PuebloPaleta, 0)
myStarters.append(Squirtle)

#     Pueblo Paleta
Krabby = pokemon('krabby',Agua, PuebloPaleta, 20)
myPokes.append(Krabby)
Horsea = pokemon('Horsea',Agua, PuebloPaleta, 60)
myPokes.append(Horsea)
Magikarp = pokemon('Magikarp', Agua, PuebloPaleta, 20)
myPokes.append(Magikarp)
Psyduck = pokemon('Psyduck', Agua, PuebloPaleta,1)
myPokes.append(Psyduck)
Shellder = pokemon('Shellder', Agua, PuebloPaleta, 40)
myPokes.append(Shellder)

#   Ciudad Verde

Poliwag = pokemon('Poliwag', Agua, CiudadVerde, 60)
myPokes.append(Poliwag)
Goldeen = pokemon('Goldeen', Agua, CiudadVerde, 20)
myPokes.append(Goldeen)
Magikarp = pokemon('Magikarp', Agua, CiudadVerde,100)
myPokes.append(Magikarp)
Psyduck = pokemon('Psyduck', Agua, CiudadVerde, 5)
myPokes.append(Psyduck)
Poliwhirl = pokemon('Poliwhirl', Agua, CiudadVerde, 40)
myPokes.append(Poliwhirl)
Slowpoke = pokemon('Slowpoke', Psiquico, CiudadVerde, 5)
myPokes.append(Slowpoke)
Gyarados = pokemon('Gyarados', Agua, CiudadVerde, 15)
myPokes.append(Gyarados)
Pidgey = pokemon('Pidgey', Volador, CiudadVerde, 45)
myPokes.append(Pidgey)
Rattata = pokemon('Rattata', Normal, CiudadVerde, 45)
myPokes.append(Rattata)

#   bosque verde

Caterpie = pokemon('Caterpie',Bicho ,BosqueVerde,40)
myPokes.append(Caterpie)
Metapod = pokemon('Metapod',Bicho , BosqueVerde, 5)
myPokes.append(Metapod)
Weedle = pokemon('Weedle', Veneno ,BosqueVerde, 40)
myPokes.append(Weedle)
Kakuna = pokemon('Kakuna', Bicho , BosqueVerde, 10)
myPokes.append(Kakuna)
Pikachu = pokemon('Pikachu', Electrico, BosqueVerde, 5)
myPokes.append(Pikachu)

#   ciudad plateada

Caterpie = pokemon('Caterpie',Bicho ,CiudadPlateada,40)
myPokes.append(Caterpie)
Metapod = pokemon('Metapod',Bicho , CiudadPlateada, 5)
myPokes.append(Metapod)
Weedle = pokemon('Weedle', Veneno ,CiudadPlateada, 40)
myPokes.append(Weedle)
Kakuna = pokemon('Kakuna', Bicho , CiudadPlateada, 10)
myPokes.append(Kakuna)
Pikachu = pokemon('Pikachu', Electrico, CiudadPlateada, 5)
Pidgey = pokemon('Pidgey', Volador, CiudadPlateada, 30)
myPokes.append(Pidgey)
Spearow = pokemon('Spearow', Volador, CiudadPlateada, 35)
myPokes.append(Spearow)
NidoranF = pokemon('NidoranF', Veneno, CiudadPlateada, 1)
myPokes.append(NidoranF)
NidoranM = pokemon('NidoranM', Veneno, CiudadPlateada, 14)
myPokes.append(NidoranM)
Jigglypuff = pokemon('Jigglypuff', Normal, CiudadPlateada, 10)
myPokes.append(Jigglypuff)
Mankey = pokemon('Mankey', Lucha, CiudadPlateada,10)
myPokes.append(Mankey)
Sandshrew = pokemon('Sandshrew', Tierra, CiudadPlateada, 25)
myPokes.append(Sandshrew)
Ekans = pokemon('Ekans', Veneno, CiudadPlateada, 25)
myPokes.append(Ekans)
Clefairy = pokemon('Clefairy', Normal, CiudadPlateada, 1)
myPokes.append(Clefairy)
Zubat = pokemon('Zubat', Veneno, CiudadPlateada, 69)
myPokes.append(Zubat)
Paras = pokemon('Paras', Hierba, CiudadPlateada, 5)
myPokes.append(Paras)
Geodude = pokemon('Geodude', Roca, CiudadPlateada, 25)
myPokes.append(Geodude)

#   monte luna

Pidgey = pokemon('Pidgey', Volador, MonteLuna, 30)
myPokes.append(Pidgey)
Spearow = pokemon('Spearow', Volador, MonteLuna, 35)
myPokes.append(Spearow)
NidoranF = pokemon('NidoranF', Veneno, MonteLuna, 1)
myPokes.append(NidoranF)
NidoranM = pokemon('NidoranM', Veneno, MonteLuna, 14)
myPokes.append(NidoranM)
Jigglypuff = pokemon('Jigglypuff', Normal, MonteLuna, 10)
myPokes.append(Jigglypuff)
Mankey = pokemon('Mankey', Lucha, MonteLuna,10)
myPokes.append(Mankey)
Sandshrew = pokemon('Sandshrew', Tierra, MonteLuna, 25)
myPokes.append(Sandshrew)
Ekans = pokemon('Ekans', Veneno, MonteLuna, 25)
myPokes.append(Ekans)
Clefairy = pokemon('Clefairy', Normal, MonteLuna, 1)
myPokes.append(Clefairy)
Zubat = pokemon('Zubat', Veneno, MonteLuna, 69)
myPokes.append(Zubat)
Paras = pokemon('Paras', Hierba, MonteLuna, 5)
myPokes.append(Paras)
Geodude = pokemon('Geodude', Roca, MonteLuna, 25)
myPokes.append(Geodude)

#   ciudad celeste

Krabby = pokemon('krabby',Agua, CiudadCeleste, 20)
myPokes.append(Krabby)
Horsea = pokemon('Horsea',Agua, CiudadCeleste, 60)
myPokes.append(Horsea)
Jynx = pokemon('Jynx', Psiquico, CiudadCeleste, 1)
myPokes.append(Jynx)
Mankey = pokemon('Mankey', Lucha, CiudadCeleste,5)
myPokes.append(Mankey)
Sandshrew = pokemon('Sandshrew', Tierra, CiudadCeleste, 25)
myPokes.append(Sandshrew)
Golduck = pokemon('Golduck', Agua, CiudadCeleste, 35)
myPokes.append(Golduck)
Slowpoke = pokemon('Slowpoke', Psiquico, CiudadCeleste, 65)
myPokes.append(Slowpoke)
Slowbro = pokemon('Slowbro', Psiquico, CiudadCeleste, 35)
myPokes.append(Slowbro)
Golbat = pokemon('Golbat', Veneno, CiudadCeleste,14)
myPokes.append(Golbat)
Parasect = pokemon('Parasect', Bicho, CiudadCeleste, 25)
myPokes.append(Parasect)
Primeape = pokemon('Primeape', Lucha, CiudadCeleste, 11)
myPokes.append(Primeape)
Machoke = pokemon('Machoke', Lucha, CiudadCeleste, 10)
myPokes.append(Machoke)
Magneton = pokemon('Magneton', Electrico, CiudadCeleste, 20)
myPokes.append(Magneton)
electrode = pokemon('Electrode', Electrico, CiudadCeleste, 5)
myPokes.append(electrode)
Ditto = pokemon('Ditto', Normal ,CiudadCeleste,11)
myPokes.append(Ditto)
Wobbuffet = pokemon('Wobbuffet', Psiquico, CiudadCeleste, 4)
myPokes.append(Wobbuffet)
Graveler = pokemon('Graveler', Roca, CiudadCeleste, 35)
myPokes.append(Graveler)
Kadabra = pokemon('Kadabra', Psiquico, CiudadCeleste,20)
myPokes.append(Kadabra)


#   tunel roca

Rattata = pokemon('Rattata', Normal, TunelRoca, 40)
myPokes.append(Rattata)
Spearow = pokemon('Spearow', Volador, TunelRoca , 35)
myPokes.append(Spearow)
Ekans = pokemon('Ekans', Veneno, TunelRoca, 25)
myPokes.append(Ekans)
Sandshrew = pokemon('Sandshrew', Tierra, TunelRoca , 25)
myPokes.append(Sandshrew)
Graveler = pokemon('Graveler', Roca, TunelRoca, 15)
myPokes.append(Graveler)
Zubat = pokemon('Zubat', Veneno, TunelRoca, 30)
myPokes.append(Zubat)
Onix = pokemon('Onix', Roca, TunelRoca, 10)
myPokes.append(Onix)

#   pueblo lavanda

Voltorb = pokemon('Voltorb', Electrico, PuebloLavanda, 30)
myPokes.append(Voltorb)
electrode = pokemon('Electrode', Electrico, PuebloLavanda, 1)
myPokes.append(electrode)
Magneton = pokemon('Magneton', Electrico, PuebloLavanda, 10)
myPokes.append(Magneton)
Electabuzz = pokemon('Electabuzz', Electrico, PuebloLavanda, 5)
myPokes.append(Electabuzz)
Gastly = pokemon('Gastly', Fantasma, PuebloLavanda, 90)
myPokes.append(Gastly)
Haunter = pokemon('Haunter',Fantasma, PuebloLavanda, 5)
myPokes.append(Haunter)
Cubone = pokemon('Cubone', Tierra, PuebloLavanda, 9)
myPokes.append(Cubone)

#   ciudad azafran

Pidgey = pokemon('Pidgey', Volador, CiudadAzafran, 30)
myPokes.append(Pidgey)
Ekans = pokemon('Ekans', Veneno, CiudadAzafran, 20)
myPokes.append(Ekans)
Sandshrew = pokemon('Sandshrew', Tierra, CiudadAzafran, 20)
myPokes.append(Sandshrew)
Vulpix = pokemon('Vulpix', Fuego, CiudadAzafran, 20)
myPokes.append(Vulpix)
Meowth = pokemon('Meowth', Normal, CiudadAzafran, 30)
myPokes.append(Meowth)
Growlithe = pokemon('Growlithe', Fuego, CiudadAzafran, 20)
myPokes.append(Growlithe)
Oddish = pokemon('Odish', Hierba,  CiudadAzafran, 20)
myPokes.append(Oddish)
Bellsprout = pokemon('Bellsprout', Hierba, CiudadAzafran, 10)
myPokes.append(Bellsprout)

#   ciudad azulona

Pidgey = pokemon('Pidgey', Volador, CiudadAzulona, 30)
myPokes.append(Pidgey)
Vulpix = pokemon('Vulpix', Fuego, CiudadAzulona, 10)
myPokes.append(Vulpix)
Oddish = pokemon('Odish', Hierba,  CiudadAzulona, 20)
myPokes.append(Oddish)
Meowth = pokemon('Meowth', Normal, CiudadAzulona, 40)
myPokes.append(Meowth)
Bellsprout = pokemon('Bellsprout', Hierba, CiudadAzulona, 20)
myPokes.append(Bellsprout)
Raticate = pokemon('Raticate', Normal, CiudadAzulona, 5)
myPokes.append(Raticate)
Doduo = pokemon('Doduo', Volador, CiudadAzulona, 35)
myPokes.append(Doduo)
Fearow = pokemon('Fearow', Volador, CiudadAzulona, 5)
myPokes.append(Fearow)

#   ciudad fucsia

Krabby = pokemon('krabby',Agua, CiudadFucsia, 20)
myPokes.append(Krabby)
Horsea = pokemon('Horsea',Agua, CiudadFucsia, 60)
myPokes.append(Horsea)
Magikarp = pokemon('Magikarp', Agua, CiudadFucsia, 100)
myPokes.append(Magikarp)
Psyduck = pokemon('Psyduck', Agua, CiudadFucsia,1)
myPokes.append(Psyduck)
Staryu = pokemon('Staryu', Agua, CiudadFucsia, 40)
myPokes.append(Staryu)
Bellsprout = pokemon('Bellsprout', Hierba, CiudadFucsia, 20)
myPokes.append(Bellsprout)
Oddish = pokemon('Odish', Hierba,  CiudadFucsia, 20)
myPokes.append(Oddish)
Poliwhirl = pokemon('Poliwhirl', Agua, CiudadFucsia, 40)
myPokes.append(Poliwhirl)
Meowth = pokemon('Meowth', Normal, CiudadFucsia, 40)
myPokes.append(Meowth)
Voltorb = pokemon('Voltorb', Electrico, CiudadFucsia, 30)
myPokes.append(Voltorb)
electrode = pokemon('Electrode', Electrico, CiudadFucsia, 1)
myPokes.append(electrode)
Magneton = pokemon('Magneton', Electrico, CiudadFucsia, 10)
myPokes.append(Magneton)
Electabuzz = pokemon('Electabuzz', Electrico, CiudadFucsia, 5)
myPokes.append(Electabuzz)
Gastly = pokemon('Gastly', Fantasma, CiudadFucsia, 90)
myPokes.append(Gastly)
Haunter = pokemon('Haunter',Fantasma, CiudadFucsia, 5)
myPokes.append(Haunter)
Cubone = pokemon('Cubone', Tierra, CiudadFucsia, 9)
myPokes.append(Cubone)


#   ciudad carmin

Krabby = pokemon('krabby',Agua, CiudadCarmin, 20)
myPokes.append(Krabby)
Horsea = pokemon('Horsea',Agua, CiudadCarmin, 60)
myPokes.append(Horsea)
Magikarp = pokemon('Magikarp', Agua, CiudadCarmin, 100)
myPokes.append(Magikarp)
Psyduck = pokemon('Psyduck', Agua, CiudadCarmin,1)
myPokes.append(Psyduck)
Staryu = pokemon('Staryu', Agua, CiudadCarmin, 40)
myPokes.append(Staryu)
Bellsprout = pokemon('Bellsprout', Hierba, CiudadCarmin, 20)
myPokes.append(Bellsprout)
Oddish = pokemon('Odish', Hierba,  CiudadCarmin, 20)
myPokes.append(Oddish)
Poliwhirl = pokemon('Poliwhirl', Agua, CiudadCarmin, 40)
myPokes.append(Poliwhirl)
Meowth = pokemon('Meowth', Normal, CiudadCarmin, 40)
myPokes.append(Meowth)

#   islas espuma

Krabby = pokemon('krabby',Agua, IslasEspuma, 20)
myPokes.append(Krabby)
Horsea = pokemon('Horsea',Agua, IslasEspuma, 60)
myPokes.append(Horsea)
Magikarp = pokemon('Magikarp', Agua, IslasEspuma, 20)
myPokes.append(Magikarp)
Psyduck = pokemon('Psyduck', Agua, IslasEspuma,1)
myPokes.append(Psyduck)
Shellder = pokemon('Shellder', Agua, IslasEspuma, 40)
myPokes.append(Shellder)
Pidgey = pokemon('Pidgey', Volador, IslasEspuma, 30)
myPokes.append(Pidgey)
Pidgeotto = pokemon('Pidgeotto', Volador, IslasEspuma, 5)
myPokes.append(Pidgeotto)
Gloom = pokemon('Gloom', Veneno, IslasEspuma, 5)
myPokes.append(Gloom)
Venonat = pokemon('Venonat', Veneno, IslasEspuma,30)
myPokes.append(Venonat)
Weepinbell = pokemon('Weepinbell', Hierba, IslasEspuma, 5)
myPokes.append(Weepinbell)
Doduo = pokemon('Doduo', Volador, IslasEspuma, 35)
myPokes.append(Doduo)
Fearow = pokemon('Fearow', Volador, IslasEspuma, 5)
myPokes.append(Fearow)

#   isla canela

Krabby = pokemon('krabby',Agua, IslaCanela, 20)
myPokes.append(Krabby)
Horsea = pokemon('Horsea',Agua, IslaCanela, 60)
myPokes.append(Horsea)
Magikarp = pokemon('Magikarp', Agua, IslaCanela, 20)
myPokes.append(Magikarp)
Psyduck = pokemon('Psyduck', Agua, IslaCanela,1)
myPokes.append(Psyduck)
Shellder = pokemon('Shellder', Agua, IslaCanela, 40)
myPokes.append(Shellder)
Gyarados = pokemon('Gyarados', Agua, IslaCanela, 15)
myPokes.append(Gyarados)
Kingler = pokemon('Kingler', Agua, IslaCanela,4)
myPokes.append(Kingler)
Seadra = pokemon('Seadra', Agua, IslaCanela, 4)
myPokes.append(Seadra)
Slowbro = pokemon('Slowbro', Psiquico, IslaCanela, 35)
myPokes.append(Slowbro)
Seel = pokemon('Seel', Hielo, IslaCanela, 1)
myPokes.append(Seel)
Rattata = pokemon('Rattata', Normal, IslaCanela, 40)
myPokes.append(Rattata)
Raticate = pokemon('Raticate', Normal, IslaCanela, 5)
myPokes.append(Raticate)
Vulpix = pokemon('Vulpix', Fuego, IslaCanela, 20)
myPokes.append(Vulpix)
Growlithe = pokemon('Growlithe', Fuego, IslaCanela, 20)
myPokes.append(Growlithe)
Grimer = pokemon('Grimer', Veneno, IslaCanela, 10)
myPokes.append(Grimer)
Muk = pokemon('Muk', Veneno, IslaCanela, 5)
myPokes.append(Muk)
Koffing = pokemon('Koffing', Veneno, IslaCanela,30)
myPokes.append(Koffing)
Weezing = pokemon('Weezing', Veneno, IslaCanela, 5)
myPokes.append(Weezing)

def pokAgainstGym(city):
  possible = []
  possible2 = []
  for i in myPokes:
    if i.site == city:
      possible.append(i)
  for i in possible:
    if i.ptype.name in city.pokeType.debilidades:
      possible2.append(i)
  myRand = random.randint(0, len(possible2)-1)#indice dentro de la lista de possible 2(coinciden en tipo y localizacion)
  found = False
  rate = 0
  while not found:
    rate = random.randint(1,100)
    if rate>0 and rate<=possible2[myRand].rate:
      return possible2[myRand]
    else:
      myRand = random.randint(0, len(possible2)-1)


def juegoCompleto():

  trayectos = ['pricipal.png','PuebloPaleta_CuidadPlateada.png', 'CuidadPalteada_CuidadCeleste.png', 'CuidadCeleste_CuidadCarmin.png', 'CuidadCarmin_CuidadAzulona.png', 'CuidadAzulona_CuidadFucsia.png','CuidadFucsia_CuidadAzafran.png','CuidadAzafran_IslaCanela.png','IslaCanela_CuidadVerde.png']
  document = Document()
  document.add_heading('------RESULTADOS FINALES------', 0)
  par1 = document.add_paragraph('Tenga en cuenta el mapa inicial de la región de Kanto: ')
  document.add_picture('principal.png', width=Inches(5.25))
  par2 = document.add_paragraph('A continuación se mostrarán los resultados del programa que permite hallar los caminos mínimos entre cada gimnasio, y a su vez el pokemón que se debe atrapar en este trayecto.')
  i=1
  for g in listaGimnasiosEnOrden:
      p = pokAgainstGym(g)
      par3 = document.add_paragraph('El gimnasio número '+str(i)+' es el de '+str(g.name)+' de tipo '+str(g.pokeType.name)+'. De los pokemón que se encuentran en la zona el mas útil para vencer el gimnasio es '+str(p.name)+' de tipo '+str(p.ptype.name)+'.\n')
      document.add_picture(trayectos[i], width=Inches(5.25))#foto trayecto
      name = str(p.name) + '.png'
      document.add_picture(name, width=Inches(2.10))#imagen pokemon
      i+=1
  document.add_page_break()
  document.save('ResultadosFinales.docx')

#FIN

juegoCompleto()
